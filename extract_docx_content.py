from docx import Document
import os

def extract_text_images_docx(docx_path, output_dir):
    document = Document(docx_path)
    text_file_path = os.path.join(output_dir, "docx_extracted_text.txt")
    images_dir = os.path.join(output_dir, "docx_images")
    os.makedirs(images_dir, exist_ok=True)

    def get_font_details(paragraph):
        fonts = set()
        sizes = set()
        colors = set()
        bold = False
        italic = False
        for run in paragraph.runs:
            fonts.add(run.font.name)
            if run.font.size:
                sizes.add(run.font.size.pt)
            if run.font.color and run.font.color.rgb:
                colors.add(run.font.color.rgb)
            if run.bold:
                bold = True
            if run.italic:
                italic = True
        return {
            "font": list(fonts)[0] if len(fonts) == 1 else "Mixed",
            "size": list(sizes)[0] if len(sizes) == 1 else "Mixed",
            "color": list(colors)[0] if len(colors) == 1 else "Mixed",
            "bold": bold,
            "italic": italic
        }

    with open(text_file_path, "w", encoding="utf-8") as text_file:
        for paragraph in document.paragraphs:
            paragraph_text = paragraph.text.strip()
            if paragraph_text:
                font_details = get_font_details(paragraph)
                text_file.write(f"{{Text: {paragraph_text}, Font: {font_details['font']}, Size: {font_details['size']}, Bold: {font_details['bold']}, Italic: {font_details['italic']}, Color: {font_details['color']}}}\n\n")

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = " ".join([p.text.strip() for p in cell.paragraphs])
                    if cell_text:
                        font_details = get_font_details(cell.paragraphs[0])  # Assuming uniform styling in cells
                        text_file.write(f"{{Text: {cell_text}, Font: {font_details['font']}, Size: {font_details['size']}, Bold: {font_details['bold']}, Italic: {font_details['italic']}, Color: {font_details['color']}}}\t")
                text_file.write("\n")
            text_file.write("\n")

    for rel in document.part.rels.values():
        if "image" in rel.target_ref:
            img = rel.target_part
            img_bytes = img.blob
            img_ext = img.content_type.split('/')[-1]
            img_filename = os.path.join(images_dir, f"image{len(os.listdir(images_dir)) + 1}.{img_ext}")
            with open(img_filename, "wb") as img_file:
                img_file.write(img_bytes)

if __name__ == "__main__":
    docx_path = "input_files/docx_mock_file.docx"
    output_dir = "output_files"
    extract_text_images_docx(docx_path, output_dir)
