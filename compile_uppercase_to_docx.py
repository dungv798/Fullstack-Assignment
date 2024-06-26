from docx import Document
from docx.shared import Pt, RGBColor
import os


def compile_uppercase_to_docx(input_text_file, output_docx_file):
    document = Document()

    with open(input_text_file, "r", encoding="utf-8") as file:
        lines = file.readlines()

    for line in lines:
        line = line.strip()
        if line.startswith("Page "):
            # Add page number as a heading
            document.add_heading(line, level=1)
        elif line.startswith("{'Text': "):
            try:
                text_attributes = eval(line)
                p = document.add_paragraph()
                run = p.add_run(text_attributes["Text"].upper())
                run.font.name = text_attributes["Font"]
                run.font.size = Pt(text_attributes["Size"])
                run.bold = text_attributes["Bold"]
                run.italic = text_attributes["Italic"]
                run.font.color.rgb = RGBColor((text_attributes["Color"] >> 16) & 0xFF,
                                              (text_attributes["Color"] >> 8) & 0xFF, (text_attributes["Color"]) & 0xFF)
            except Exception as e:
                print(f"Error processing line: {line} - {e}")

    document.save(output_docx_file)


# Example usage
if __name__ == "__main__":
    input_text_file = "output_files/pdf_extracted_text.txt"
    output_docx_file = "output_files/uppercase_docx_output.docx"
    compile_uppercase_to_docx(input_text_file, output_docx_file)
